#!/usr/bin/env python3
"""
Generate DOCX performance reports from a scoring-report.json.

Reads the JSON output from grade-pset.py and produces one .docx per student
with score summaries, cell-referenced error explanations, and pattern-based
error grouping.

Usage:
    python3 generate-pset-report.py --input scoring-report.json --outdir reports/

    python3 generate-pset-report.py --input scoring-report.json --outdir reports/ \\
        --student "Cole Barnett"
"""

import argparse
import json
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Pattern descriptions — human-readable explanations for each error pattern
# ---------------------------------------------------------------------------

PATTERN_DESCRIPTIONS = {
    "missing_answer": (
        "The cell is blank or contains zero. Make sure every graded item has "
        "a formula or value entered."
    ),
    "sign_error": (
        "Your answer has the opposite sign of the correct value. Check whether "
        "you need to subtract in the other direction or remove an extra negative sign."
    ),
    "gave_monthly_not_annual": (
        "You appear to have reported a monthly value where an annualized value "
        "was expected. To annualize, compound the monthly return: "
        "(1 + r_monthly)^12 - 1, or multiply the monthly standard deviation by SQRT(12)."
    ),
    "annualization_error": (
        "You multiplied a monthly value by 12 to annualize. For returns, the correct "
        "approach is to compound: (1 + r_monthly)^12 - 1. Simple multiplication "
        "overstates the result due to the compounding effect."
    ),
    "used_population_stat": (
        "Your formula uses a population statistic (STDEV.P or VAR.P) instead of the "
        "sample statistic (STDEV.S or VAR.S). When working with a sample of historical "
        "returns, use the sample versions which divide by (n-1)."
    ),
    "wrong_range": (
        "Your formula references a data range that appears to be truncated or "
        "does not cover the full dataset. Double-check that your cell range "
        "extends from the first data row to the last."
    ),
    "hardcoded": (
        "This cell contains a hardcoded number instead of a formula. Use formulas "
        "that reference your data so your answers update automatically if inputs change."
    ),
    "wrong_item_value": (
        "Your answer matches the correct value for a different item. You may have "
        "referenced the wrong cell or used the wrong input in your calculation."
    ),
}

# Patterns that represent a common conceptual error (group these together)
CONCEPTUAL_PATTERNS = {
    "annualization_error",
    "gave_monthly_not_annual",
    "used_population_stat",
    "wrong_range",
    "sign_error",
}


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _pct(numerator, denominator):
    """Format as percentage string."""
    if denominator == 0:
        return "N/A"
    return f"{numerator / denominator * 100:.0f}%"


def _cell_ref(sheet, cell):
    """Build a readable cell reference like 'Q1 → cell N14'."""
    if cell:
        return f"{sheet} → cell {cell}"
    return f"{sheet}"


def _fmt_value(v):
    """Format a numeric value for display."""
    if v is None:
        return "(blank)"
    try:
        v = float(v)
    except (ValueError, TypeError):
        return str(v)
    if abs(v) >= 100:
        return f"{v:,.2f}"
    if abs(v) >= 1:
        return f"{v:.4f}"
    return f"{v:.6f}"


def _sanitize_filename(name):
    """Convert a student name to a safe filename."""
    return re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "-").lower()


def _tone_opening(pct):
    """Return a tone-appropriate opening sentence based on score percentage."""
    if pct >= 0.90:
        return "Excellent work on this assignment."
    elif pct >= 0.80:
        return "Good work overall."
    elif pct >= 0.70:
        return "Solid effort — review the items below to strengthen your understanding."
    elif pct >= 0.60:
        return "Review the errors below carefully — there are several areas to work on."
    else:
        return "This assignment needs significant improvement. Study the errors below."


def _tone_closing(pct, top_patterns):
    """Return a tone-appropriate closing paragraph."""
    if pct >= 0.90:
        parts = ["Strong performance."]
        if top_patterns:
            parts.append(f"The main area to refine is {top_patterns[0].lower().replace('_', ' ')}.")
        parts.append("Keep up the careful work on the next assignment.")
        return " ".join(parts)
    elif pct >= 0.70:
        focus = ""
        if top_patterns:
            focus = f" Focus especially on {' and '.join(p.lower().replace('_', ' ') for p in top_patterns[:2])}."
        return f"You have a solid foundation.{focus} Reviewing these concepts before the next assignment will help you improve."
    else:
        focus = ""
        if top_patterns:
            focus = f" The most impactful areas to study are {' and '.join(p.lower().replace('_', ' ') for p in top_patterns[:2])}."
        return f"Spend time reviewing the fundamentals covered in this assignment.{focus} If you have questions, come to office hours."


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _add_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def _add_score_table(doc, student, report):
    """Add the score summary table."""
    # Build rows from per-sheet scores
    rows = []
    sheets_order = list(report.get("sheets_summary", {}).keys())
    # Fall back to whatever sheets are in the student scores
    if not sheets_order:
        sheets_order = list(student["scores"].keys())

    for sheet in sheets_order:
        if sheet in student["scores"]:
            s = student["scores"][sheet]
            rows.append((sheet, s["right"], s["total"], _pct(s["right"], s["total"])))

    rows.append(("Total", student["total_numerical"], student["total_possible"],
                 _pct(student["total_numerical"], student["total_possible"])))

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(["Section", "Correct", "Total", "Score"]):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for r_idx, (section, right, total, pct_str) in enumerate(rows, 1):
        table.rows[r_idx].cells[0].text = section
        table.rows[r_idx].cells[1].text = str(right)
        table.rows[r_idx].cells[2].text = str(total)
        table.rows[r_idx].cells[3].text = pct_str
        for c_idx in range(1, 4):
            for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bold the total row
        if r_idx == len(rows):
            for c_idx in range(4):
                for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _group_errors(items):
    """Group wrong items by pattern for conceptual grouping.

    Returns:
        list of (group_label, pattern_key, items_list) tuples,
        ordered by number of items descending.
    """
    groups = {}
    ungrouped = []

    for item in items:
        pattern = item.get("pattern_match")
        if pattern and pattern in CONCEPTUAL_PATTERNS:
            groups.setdefault(pattern, []).append(item)
        elif pattern and pattern in PATTERN_DESCRIPTIONS:
            groups.setdefault(pattern, []).append(item)
        else:
            ungrouped.append(item)

    # Sort groups by count descending
    sorted_groups = sorted(groups.items(), key=lambda x: -len(x[1]))

    result = []
    for pattern, group_items in sorted_groups:
        label = pattern.replace("_", " ").title()
        result.append((label, pattern, group_items))

    if ungrouped:
        result.append(("Other Errors", None, ungrouped))

    return result


def _explain_item(item):
    """Generate a narrative explanation for a wrong item.

    Returns a list of (text, is_code) tuples for building a paragraph
    with inline code formatting where formulas are cited.
    """
    sheet = item.get("sheet", "")
    cell = item.get("student_cell")
    question = item.get("question", "")
    formula = item.get("student_formula")
    student_val = item.get("student_value")
    pattern = item.get("pattern_match")

    # Build cell location string
    loc = f"cell {cell} on {sheet}" if cell else f"on sheet {sheet}"

    # Lowercase the question for embedding in sentences, preserving proper nouns
    if question:
        # Only lowercase the first character if it's not part of a proper noun / acronym
        first_word = question.split()[0] if question.split() else ""
        if first_word == first_word.upper() and len(first_word) > 1:
            # Acronym or all-caps word — keep as-is
            q_lower = question
        elif first_word[0].isupper() and first_word[1:].islower() and first_word not in (
            "Sharpe", "CAPM", "VaR", "TSLA", "WMT", "BKE", "NNI", "VMI",
        ):
            q_lower = question[0].lower() + question[1:]
        else:
            q_lower = question
    else:
        q_lower = "this item"
    # Strip trailing question marks/periods for embedding
    q_lower = q_lower.rstrip("?.").strip()

    parts = []  # list of (text, is_code) tuples

    if pattern == "missing_answer":
        parts.append((f"In {loc}, the cell is blank. You need to enter a formula for {q_lower}.", False))

    elif pattern == "sign_error":
        parts.append((f"In {loc}, your answer for {q_lower} has the wrong sign.", False))
        if formula and formula.startswith("="):
            parts.append((" Check your formula ", False))
            parts.append((formula, True))
            parts.append((" — you may need to reverse the order of subtraction or remove an extra negative.", False))
        else:
            parts.append((" Check whether you need to subtract in the other direction.", False))

    elif pattern == "gave_monthly_not_annual":
        parts.append((f"In {loc}, you reported a monthly value instead of an annualized one for {q_lower}.", False))
        if formula and formula.startswith("="):
            parts.append((" Your formula ", False))
            parts.append((formula, True))
            parts.append((" computes the monthly figure. To annualize returns, compound with (1 + r_monthly)^12 − 1. For volatility, multiply by SQRT(12).", False))
        else:
            parts.append((" To annualize returns, compound with (1 + r_monthly)^12 − 1. For volatility, multiply by SQRT(12).", False))

    elif pattern == "annualization_error":
        parts.append((f"In {loc}, your formula for {q_lower} multiplies by 12 instead of compounding.", False))
        if formula and formula.startswith("="):
            parts.append((" Your formula ", False))
            parts.append((formula, True))
            parts.append((" uses *12. The correct approach is (1 + r_monthly)^12 − 1, which accounts for the compounding effect.", False))
        else:
            parts.append((" Use (1 + r_monthly)^12 − 1 instead of multiplying by 12.", False))

    elif pattern == "used_population_stat":
        parts.append((f"In {loc}, your formula for {q_lower} uses a population statistic instead of a sample statistic.", False))
        if formula and formula.startswith("="):
            parts.append((" Your formula ", False))
            parts.append((formula, True))
            # Detect which function to suggest replacing
            formula_upper = formula.upper()
            if "STDEV.P" in formula_upper or "STDEVP" in formula_upper:
                parts.append((" uses STDEV.P — replace it with STDEV.S, which divides by (n − 1) as appropriate for sample data.", False))
            elif "VAR.P" in formula_upper or "VARP" in formula_upper:
                parts.append((" uses VAR.P — replace it with VAR.S, which divides by (n − 1) as appropriate for sample data.", False))
            else:
                parts.append((" uses a population function. Use the sample version (ending in .S) instead.", False))
        else:
            parts.append((" Use STDEV.S or VAR.S instead of STDEV.P or VAR.P when working with sample data.", False))

    elif pattern == "wrong_range":
        parts.append((f"In {loc}, the data range in your formula for {q_lower} appears to be truncated.", False))
        if formula and formula.startswith("="):
            parts.append((" Your formula ", False))
            parts.append((formula, True))
            parts.append((" does not cover the full dataset. Make sure your range extends from the first data row to the last.", False))
        else:
            parts.append((" Double-check that your cell range covers all the data rows.", False))

    elif pattern == "hardcoded":
        parts.append((f"In {loc}, you entered a hardcoded number instead of a formula for {q_lower}.", False))
        parts.append((" Use a formula that references your data so the answer updates automatically if inputs change.", False))

    elif pattern == "wrong_item_value":
        parts.append((f"In {loc}, your answer for {q_lower} appears to match a different item's value.", False))
        if formula and formula.startswith("="):
            parts.append((" Check the cell references in your formula ", False))
            parts.append((formula, True))
            parts.append((" — you may be pulling from the wrong row or column.", False))
        else:
            parts.append((" Check that you are referencing the correct inputs for this particular calculation.", False))

    else:
        # No pattern — generic explanation with formula if available
        parts.append((f"In {loc}, your answer for {q_lower} is incorrect.", False))
        if formula and formula.startswith("="):
            parts.append((" Your formula is ", False))
            parts.append((formula, True))
            parts.append((" — review the inputs and logic of this calculation.", False))

    return parts


def _add_error_group(doc, group_label, pattern_key, items, group_num):
    """Add one error group section to the document."""
    pts = len(items)
    heading_text = f"{group_num}. {group_label} (−{pts} pt{'s' if pts != 1 else ''})"
    _add_heading(doc, heading_text, level=2)

    # Group-level explanation
    if pattern_key and pattern_key in PATTERN_DESCRIPTIONS:
        p = doc.add_paragraph()
        p.style = "Body Text"
        run = p.add_run(PATTERN_DESCRIPTIONS[pattern_key])
        run.italic = True

    # For missing_answer with many items, consolidate into a summary
    if pattern_key == "missing_answer" and len(items) > 5:
        # Group by sheet
        by_sheet = {}
        for item in items:
            sheet = item.get("sheet", "Unknown")
            by_sheet.setdefault(sheet, []).append(item)

        for sheet, sheet_items in by_sheet.items():
            # List cells if available, otherwise count
            cells_with_ref = sorted(set(
                it.get("student_cell") for it in sheet_items if it.get("student_cell")
            ))
            p = doc.add_paragraph(style="List Bullet")
            if cells_with_ref:
                cell_list = ", ".join(cells_with_ref)
                p.add_run(f"{sheet}: {len(sheet_items)} blank items (cells {cell_list}).")
            else:
                p.add_run(f"{sheet}: {len(sheet_items)} items are blank — the entire section is unanswered.")
        return

    # Deduplicate items that map to the same cell (two grading items → same student cell)
    seen_cells = set()
    deduped_items = []
    for item in items:
        cell_key = f"{item.get('sheet')}!{item.get('student_cell')}"
        if cell_key in seen_cells and item.get("student_cell"):
            continue  # skip duplicate cell
        seen_cells.add(cell_key)
        deduped_items.append(item)

    # Each item as a narrative bullet
    for item in deduped_items:
        explanation_parts = _explain_item(item)

        p = doc.add_paragraph(style="List Bullet")
        for text, is_code in explanation_parts:
            run = p.add_run(text)
            if is_code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)


def generate_student_report(student, report, outdir):
    """Generate a DOCX performance report for one student."""
    doc = Document()

    # -- Style tweaks --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    assignment = report["assignment"].replace("_KEY", "").replace("_", " ")
    pct = student["percentage"]
    name = student["name"]

    # -- Header --
    _add_heading(doc, f"{assignment}", level=1)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run(f"Student: {name}")
    run.bold = True
    run.font.size = Pt(13)

    score_p = doc.add_paragraph()
    run = score_p.add_run(f"Score: {student['total_numerical']} / {student['total_possible']} ({_pct(student['total_numerical'], student['total_possible'])})")
    run.bold = True
    run.font.size = Pt(14)
    if pct >= 0.90:
        run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x23)
    elif pct >= 0.70:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)
    else:
        run.font.color.rgb = RGBColor(0xB8, 0x28, 0x11)

    doc.add_paragraph(_tone_opening(pct))

    # -- Score Summary Table --
    _add_heading(doc, "Score Summary", level=1)
    _add_score_table(doc, student, report)
    doc.add_paragraph()  # spacer

    # -- Error details --
    wrong_items = student.get("items", [])
    if wrong_items:
        _add_heading(doc, "Where You Lost Points", level=1)

        error_groups = _group_errors(wrong_items)
        top_patterns = [pat for _, pat, _ in error_groups if pat]

        for i, (label, pattern_key, group_items) in enumerate(error_groups, 1):
            _add_error_group(doc, label, pattern_key, group_items, i)
    else:
        _add_heading(doc, "Where You Lost Points", level=1)
        doc.add_paragraph("No errors — perfect score!")
        top_patterns = []

    # -- Closing --
    _add_heading(doc, "Next Steps", level=1)
    doc.add_paragraph(_tone_closing(pct, top_patterns))

    # -- Save --
    safe_name = _sanitize_filename(name)
    assignment_short = _sanitize_filename(report["assignment"].replace("_KEY", ""))
    filename = f"{assignment_short}-{safe_name}.docx"
    filepath = os.path.join(outdir, filename)
    doc.save(filepath)
    return filepath


# ---------------------------------------------------------------------------
# Instructor Report
# ---------------------------------------------------------------------------

def generate_instructor_report(report, outdir):
    """Generate a class-level instructor summary report."""
    from collections import Counter
    import statistics

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    assignment = report["assignment"].replace("_KEY", "").replace("_", " ")
    students = report["students"]
    total_possible = report["grading_map"][-1]["correct_value"] if report["grading_map"] else 0  # unused
    total_possible = len(report["grading_map"])
    pcts = [s["percentage"] for s in students]

    # -- Header --
    _add_heading(doc, f"{assignment}", level=1)
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Instructor Summary Report")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # -- Class Overview --
    _add_heading(doc, "Class Overview", level=1)
    avg_pct = statistics.mean(pcts)
    median_pct = statistics.median(pcts)
    std_pct = statistics.stdev(pcts) if len(pcts) > 1 else 0

    overview_data = [
        ("Students", str(len(students))),
        ("Items graded", str(total_possible)),
        ("Average", f"{avg_pct:.0%}"),
        ("Median", f"{median_pct:.0%}"),
        ("Std Dev", f"{std_pct:.0%}"),
        ("Min", f"{min(pcts):.0%}"),
        ("Max", f"{max(pcts):.0%}"),
        ("Perfect scores", str(sum(1 for p in pcts if p >= 0.999))),
    ]
    table = doc.add_table(rows=len(overview_data), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(overview_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # -- Score Distribution --
    _add_heading(doc, "Score Distribution", level=1)
    buckets = [
        ("90-100%", 0.90, 1.01),
        ("80-89%", 0.80, 0.90),
        ("70-79%", 0.70, 0.80),
        ("60-69%", 0.60, 0.70),
        ("50-59%", 0.50, 0.60),
        ("Below 50%", 0.00, 0.50),
    ]
    dist_table = doc.add_table(rows=len(buckets) + 1, cols=4)
    dist_table.style = "Light Grid Accent 1"
    dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Range", "Count", "Percent", ""]):
        cell = dist_table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, (label, lo, hi) in enumerate(buckets, 1):
        count = sum(1 for p in pcts if lo <= p < hi)
        frac = count / len(pcts) if pcts else 0
        bar_len = int(frac * 30)
        bar = "\u2588" * bar_len

        dist_table.rows[r_idx].cells[0].text = label
        dist_table.rows[r_idx].cells[1].text = str(count)
        dist_table.rows[r_idx].cells[2].text = f"{frac:.0%}"
        dist_table.rows[r_idx].cells[3].text = bar
        for c in range(1, 3):
            dist_table.rows[r_idx].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # -- Per-Section Analysis --
    _add_heading(doc, "Per-Section Performance", level=1)
    sheets_order = list(report.get("sheets_summary", {}).keys())
    if sheets_order:
        sec_table = doc.add_table(rows=len(sheets_order) + 1, cols=5)
        sec_table.style = "Light Grid Accent 1"
        sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(["Section", "Items", "Class Avg", "Min", "Max"]):
            cell = sec_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for r_idx, sheet in enumerate(sheets_order, 1):
            info = report["sheets_summary"][sheet]
            n_items = info["total_items"]
            # Compute per-sheet stats from student data
            sheet_pcts = []
            for s in students:
                if sheet in s["scores"]:
                    sc = s["scores"][sheet]
                    if sc["total"] > 0:
                        sheet_pcts.append(sc["right"] / sc["total"])
            sec_table.rows[r_idx].cells[0].text = sheet
            sec_table.rows[r_idx].cells[1].text = str(n_items)
            sec_table.rows[r_idx].cells[2].text = f"{statistics.mean(sheet_pcts):.0%}" if sheet_pcts else "N/A"
            sec_table.rows[r_idx].cells[3].text = f"{min(sheet_pcts):.0%}" if sheet_pcts else "N/A"
            sec_table.rows[r_idx].cells[4].text = f"{max(sheet_pcts):.0%}" if sheet_pcts else "N/A"
            for c in range(1, 5):
                sec_table.rows[r_idx].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    # -- Common Error Patterns --
    _add_heading(doc, "Common Error Patterns", level=1)
    pattern_counts = Counter()
    pattern_students = Counter()  # how many students had each pattern
    for s in students:
        seen_patterns = set()
        for item in s.get("items", []):
            pat = item.get("pattern_match")
            if pat:
                pattern_counts[pat] += 1
                seen_patterns.add(pat)
        for pat in seen_patterns:
            pattern_students[pat] += 1

    total_wrong = sum(len(s.get("items", [])) for s in students)
    unmatched = total_wrong - sum(pattern_counts.values())

    if pattern_counts:
        pat_table = doc.add_table(rows=len(pattern_counts) + 2, cols=4)
        pat_table.style = "Light Grid Accent 1"
        pat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(["Pattern", "Occurrences", "Students Affected", "Description"]):
            cell = pat_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for r_idx, (pat, count) in enumerate(pattern_counts.most_common(), 1):
            pat_table.rows[r_idx].cells[0].text = pat.replace("_", " ").title()
            pat_table.rows[r_idx].cells[1].text = str(count)
            pat_table.rows[r_idx].cells[2].text = f"{pattern_students[pat]} / {len(students)}"
            desc = PATTERN_DESCRIPTIONS.get(pat, "")
            pat_table.rows[r_idx].cells[3].text = desc[:80] + ("..." if len(desc) > 80 else "")
            for c in range(1, 3):
                pat_table.rows[r_idx].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Unmatched row
        last_row = len(pattern_counts) + 1
        pat_table.rows[last_row].cells[0].text = "No pattern matched"
        pat_table.rows[last_row].cells[1].text = str(unmatched)
        pat_table.rows[last_row].cells[2].text = "—"
        pat_table.rows[last_row].cells[3].text = "Errors not matching any heuristic"
        for run in pat_table.rows[last_row].cells[0].paragraphs[0].runs:
            run.italic = True
        for c in range(1, 3):
            pat_table.rows[last_row].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # -- Most Difficult Items --
    _add_heading(doc, "Most Difficult Items", level=1)
    doc.add_paragraph("Items with the highest error rates across the class.")

    # Count wrong answers per item_id
    item_wrong_counts = Counter()
    for s in students:
        for item in s.get("items", []):
            item_wrong_counts[item["item_id"]] += 1

    # Build lookup for question text
    item_questions = {g["item_id"]: g for g in report["grading_map"]}

    top_items = item_wrong_counts.most_common(15)
    if top_items:
        diff_table = doc.add_table(rows=len(top_items) + 1, cols=5)
        diff_table.style = "Light Grid Accent 1"
        diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(["Item", "Sheet", "Wrong", "Error Rate", "Question"]):
            cell = diff_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for r_idx, (item_id, wrong_count) in enumerate(top_items, 1):
            gm = item_questions.get(item_id, {})
            error_rate = wrong_count / len(students) if students else 0
            diff_table.rows[r_idx].cells[0].text = item_id
            diff_table.rows[r_idx].cells[1].text = gm.get("sheet", "")
            diff_table.rows[r_idx].cells[2].text = str(wrong_count)
            diff_table.rows[r_idx].cells[3].text = f"{error_rate:.0%}"
            question = gm.get("question", "")
            diff_table.rows[r_idx].cells[4].text = question[:60] + ("..." if len(question) > 60 else "")
            for c in range(2, 4):
                diff_table.rows[r_idx].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # -- Student Roster --
    _add_heading(doc, "Student Roster", level=1)

    # Build header: Name, Total, Pct, then per-sheet scores
    col_headers = ["Name", "Score", "Pct"]
    for sheet in sheets_order:
        col_headers.append(sheet)

    roster = doc.add_table(rows=len(students) + 1, cols=len(col_headers))
    roster.style = "Light Grid Accent 1"
    roster.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(col_headers):
        cell = roster.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sort students by percentage descending
    sorted_students = sorted(students, key=lambda s: -s["percentage"])
    for r_idx, s in enumerate(sorted_students, 1):
        roster.rows[r_idx].cells[0].text = s["name"]
        roster.rows[r_idx].cells[1].text = f"{s['total_numerical']}/{s['total_possible']}"
        roster.rows[r_idx].cells[2].text = _pct(s["total_numerical"], s["total_possible"])
        for c_idx, sheet in enumerate(sheets_order, 3):
            if sheet in s["scores"]:
                sc = s["scores"][sheet]
                roster.rows[r_idx].cells[c_idx].text = f"{sc['right']}/{sc['total']}"
            else:
                roster.rows[r_idx].cells[c_idx].text = "—"
        for c in range(1, len(col_headers)):
            roster.rows[r_idx].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Color-code low scorers
        if s["percentage"] < 0.60:
            for c in range(len(col_headers)):
                for run in roster.rows[r_idx].cells[c].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xB8, 0x28, 0x11)

    doc.add_paragraph()

    # -- Students Needing Attention --
    at_risk = [s for s in sorted_students if s["percentage"] < 0.60]
    if at_risk:
        _add_heading(doc, f"Students Needing Attention ({len(at_risk)})", level=1)
        doc.add_paragraph("Students scoring below 60%. Consider reaching out individually.")
        for s in at_risk:
            wrong_items = s.get("items", [])
            patterns = Counter(it.get("pattern_match") for it in wrong_items if it.get("pattern_match"))
            top_pat = ", ".join(f"{p.replace('_',' ')} ({c})" for p, c in patterns.most_common(3)) if patterns else "no patterns detected"

            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{s['name']}: {s['total_numerical']}/{s['total_possible']} ({_pct(s['total_numerical'], s['total_possible'])})")
            run.bold = True
            p.add_run(f" — {len(wrong_items)} errors. Top patterns: {top_pat}.")

    # -- Save --
    assignment_short = _sanitize_filename(report["assignment"].replace("_KEY", ""))
    filename = f"{assignment_short}-instructor-summary.docx"
    filepath = os.path.join(outdir, filename)
    doc.save(filepath)
    return filepath


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate DOCX performance reports from scoring-report.json",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--input", required=True,
                        help="Path to scoring-report.json from grade-pset.py")
    parser.add_argument("--outdir", required=True,
                        help="Output directory for .docx files")
    parser.add_argument("--student",
                        help="Generate report for a single student (name match)")
    parser.add_argument("--instructor", action="store_true",
                        help="Generate instructor summary report instead of student reports")

    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    with open(args.input) as f:
        report = json.load(f)

    os.makedirs(args.outdir, exist_ok=True)

    if args.instructor:
        filepath = generate_instructor_report(report, args.outdir)
        print(f"Instructor summary written to {filepath}")
        return

    students = report["students"]
    if args.student:
        matches = [s for s in students if args.student.lower() in s["name"].lower()]
        if not matches:
            print(f"No student matching '{args.student}' found.")
            print(f"Available: {', '.join(s['name'] for s in students[:10])}...")
            sys.exit(1)
        students = matches

    print(f"Generating reports for {len(students)} student(s)...")
    for i, student in enumerate(students, 1):
        filepath = generate_student_report(student, report, args.outdir)
        wrong = len(student.get("items", []))
        print(f"  [{i}/{len(students)}] {student['name']}: "
              f"{student['total_numerical']}/{student['total_possible']} "
              f"({_pct(student['total_numerical'], student['total_possible'])}) "
              f"— {wrong} errors → {os.path.basename(filepath)}")

    print(f"\nDone. {len(students)} reports written to {args.outdir}/")


if __name__ == "__main__":
    main()
