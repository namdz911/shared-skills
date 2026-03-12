#!/usr/bin/env python3
"""
Convert markdown feedback reports to DOCX files.

Usage:
    python3 md2docx.py --indir reports/ --outdir docx/
    python3 md2docx.py --indir reports/ --outdir docx/ --file "report-name.md"
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def parse_md_to_blocks(md_text):
    """Parse markdown into a list of structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            i += 1
            continue

        # Table — collect all rows
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and lines[i].strip():
                if re.match(r"^[-*]\s+", lines[i]):
                    items.append(lines[i].strip()[2:].strip())
                elif lines[i].startswith("  ") and items:
                    # continuation line
                    items[-1] += " " + lines[i].strip()
                else:
                    break
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue

        # Bold score line (e.g., **Score: 61 / 79 (77%)**)
        if line.strip().startswith("**Score:") or line.strip().startswith("**Items affected:"):
            blocks.append({"type": "bold_line", "text": line.strip()})
            i += 1
            continue

        # Regular paragraph — collect until blank line or structural element
        para_lines = []
        while i < len(lines) and lines[i].strip():
            if re.match(r"^#{1,4}\s+", lines[i]):
                break
            if lines[i].strip().startswith("|"):
                break
            if re.match(r"^[-*]\s+\*\*", lines[i]):
                break
            if re.match(r"^-{3,}$", lines[i].strip()):
                break
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(l.strip() for l in para_lines)})

    return blocks


def add_formatted_text(paragraph, text):
    """Add text with inline markdown formatting (bold, italic, code)."""
    # Split on bold, italic, and code patterns
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table(table_lines):
    """Parse markdown table lines into header + rows."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    # Filter out separator rows (like |---|---|---|)
    data_rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r)]
    if len(data_rows) < 2:
        return None, []
    return data_rows[0], data_rows[1:]


def get_score_color(pct):
    """Return RGB color based on score percentage."""
    if pct >= 90:
        return RGBColor(0x1B, 0x7F, 0x3B)  # Green
    elif pct >= 80:
        return RGBColor(0x2E, 0x6B, 0x9E)  # Blue
    elif pct >= 70:
        return RGBColor(0xB8, 0x86, 0x0B)  # Dark goldenrod
    elif pct >= 60:
        return RGBColor(0xCC, 0x66, 0x00)  # Orange
    else:
        return RGBColor(0xCC, 0x00, 0x00)  # Red


def _add_hr_border(paragraph):
    """Add a bottom border to a paragraph to render a horizontal rule."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def md_to_docx(md_text, output_path):
    """Convert markdown feedback report to DOCX."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    blocks = parse_md_to_blocks(md_text)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            text = block["text"]
            # Strip markdown bold from headings
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

            if level == 1:
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p = doc.add_heading("", level=level)
                add_formatted_text(p, text)

        elif btype == "bold_line":
            text = block["text"]
            # Extract score percentage for coloring
            pct_match = re.search(r"\((\d+)%\)", text)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            if pct_match:
                pct = int(pct_match.group(1))
                for run in p.runs:
                    run.font.color.rgb = get_score_color(pct)
                    run.font.size = Pt(13)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            add_formatted_text(p, block["text"])

        elif btype == "bullets":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_text(p, item)

        elif btype == "table":
            header, rows = parse_table(block["lines"])
            if header is None:
                continue
            ncols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for j, cell_text in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text))
                run.bold = True
                run.font.size = Pt(10)

            # Data rows
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j >= ncols:
                        break
                    cell = table.rows[i + 1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                    run = p.add_run(clean)
                    run.font.size = Pt(10)
                    # Bold the total row
                    if "Total" in clean or "total" in clean:
                        run.bold = True

            doc.add_paragraph()  # spacing after table

        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_hr_border(p)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Convert MD feedback reports to DOCX")
    parser.add_argument("--indir", required=True, help="Directory containing .md reports")
    parser.add_argument("--outdir", required=True, help="Output directory for .docx files")
    parser.add_argument("--file", help="Convert only this specific file (basename)")
    args = parser.parse_args()

    os.makedirs(args.outdir, exist_ok=True)

    if args.file:
        files = [args.file]
    else:
        files = sorted(f for f in os.listdir(args.indir) if f.endswith(".md"))

    converted = 0
    for fname in files:
        md_path = os.path.join(args.indir, fname)
        if not os.path.exists(md_path):
            print(f"  SKIP (not found): {fname}")
            continue

        docx_name = fname.replace(".md", ".docx")
        docx_path = os.path.join(args.outdir, docx_name)

        with open(md_path, "r") as f:
            md_text = f.read()

        # Skip Professor Flags section for student-facing DOCX
        if "## Professor Flags" in md_text:
            md_text = md_text[:md_text.index("## Professor Flags")].rstrip()

        md_to_docx(md_text, docx_path)
        converted += 1
        print(f"  OK: {docx_name}")

    print(f"\nConverted {converted}/{len(files)} files to {args.outdir}/")


if __name__ == "__main__":
    main()
